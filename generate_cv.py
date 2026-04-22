#!/usr/bin/env python3
"""
DHM CV Optimisation Pipeline - Document Generator
Produces consistently formatted .docx files matching the DHM brand standard.

Brand colours:
  Deep Coral   #DC6A63  (220, 106, 99)  — primary accent
  Black        #000000  (0, 0, 0)
  White        #FFFFFF  (255, 255, 255)
  Soft Blue    #B8C0CC  (184, 192, 204)  — secondary accent
  Dark Grey    #444444  (68, 68, 68)
"""
import os
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------------------
# ASSET PATHS  (images live in an assets/ folder alongside this script)
# ---------------------------------------------------------------------------
_DIR     = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = os.path.join(_DIR, 'assets', 'DHM_App_Logo.png')
SIG_PATH  = os.path.join(_DIR, 'assets', 'DHM_Report_Signature_V3.png')

# ---------------------------------------------------------------------------
# BRAND COLOURS
# ---------------------------------------------------------------------------
CORAL     = (220, 106, 99)
BLACK     = (0, 0, 0)
WHITE     = (255, 255, 255)
SOFT_BLUE = (184, 192, 204)
DARK_GREY = (68, 68, 68)
MID_GREY  = (80, 80, 80)


# ---------------------------------------------------------------------------
# LOW-LEVEL HELPERS
# ---------------------------------------------------------------------------

def _first_name(full_name):
    """Extract and title-case the first name from cv_data['name']."""
    return full_name.strip().split()[0].title()


def add_horizontal_rule(paragraph, color='DC6A63'):
    """Bottom border — used in report headings."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_top_rule(paragraph, color='DC6A63', sz='4', space='6'):
    """Top border — used in CV section headings for premium editorial look."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), sz)       # 4 = 0.5pt hairline
    top.set(qn('w:space'), space)
    top.set(qn('w:color'), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_run_letter_spacing(run, spacing_twips=30):
    """Add character spacing to a run (1 twip = 1/20 pt; 30 ≈ 1.5pt tracking)."""
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:val'), str(spacing_twips))
    rPr.append(spacing)


def set_run_font(run, name='Calibri', size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pPr.append(shd)


def add_hyperlink(paragraph, url, text, color, size=10.5, bold=False):
    """Add a clickable hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    color_el = OxmlElement('w:color')
    color_el.set(qn('w:val'), '%02x%02x%02x' % color)
    rPr.append(color_el)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(size * 2)))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement('w:b'))
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _setup_page(doc):
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin  = section.bottom_margin = Inches(0.8)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ---------------------------------------------------------------------------
# CV SECTION HELPERS  (navy palette — CV styling, not DHM brand)
# ---------------------------------------------------------------------------

def add_cv_section_heading(doc, text):
    """CV section heading: black, 12pt, coral top-rule, letter-spaced ALL CAPS."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=12, bold=True, color=BLACK)
    add_run_letter_spacing(run, spacing_twips=30)   # ~1.5pt tracking on ALL CAPS
    add_top_rule(p, color='DC6A63', sz='4', space='6')
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    run = p.add_run(str(text) if text is not None else '')
    set_run_font(run, size=10.5)
    return p


# ---------------------------------------------------------------------------
# REPORT SECTION HELPERS  (DHM brand palette)
# ---------------------------------------------------------------------------

def add_report_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, name='Inter', size=11, bold=True, color=CORAL)
    add_horizontal_rule(p, color='DC6A63')
    return p


def add_subheading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(text)
    set_run_font(r, size=10, italic=True, color=MID_GREY)
    return p


def add_body(doc, text, bold=False, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=bold, color=BLACK if bold else DARK_GREY)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    return p


def add_stage_heading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, name='Inter', size=10.5, bold=True, color=BLACK)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_approach_bullet(doc, bold_label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(bold_label + ': ')
    set_run_font(r1, size=10.5, bold=True, color=BLACK)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=DARK_GREY)
    return p


def _render_parsed_text(paragraph, text, base_color=None, base_size=10.5):
    """Render text, converting *highlighted* spans to coral bold (no asterisks)."""
    if base_color is None:
        base_color = DARK_GREY
    parts = re.split(r'\*([^*]+)\*', str(text))
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:          # inside asterisks → coral bold
            r = paragraph.add_run(part)
            set_run_font(r, size=base_size, bold=True, color=CORAL)
        else:                    # outside asterisks → base colour, not bold
            r = paragraph.add_run(part)
            set_run_font(r, size=base_size, color=base_color)


def add_numbered_item(doc, number, bold_title, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    # Coral bold number prefix
    r1 = p.add_run(f'{number}. ')
    set_run_font(r1, size=10.5, bold=True, color=CORAL)
    # Merge title + text into one string — no forced bold, asterisks become coral
    content = str(bold_title) if bold_title else ''
    if text:
        content += (' - ' if content else '') + str(text)
    _render_parsed_text(p, content)
    return p


def add_next_step_bullet(doc, bold_label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(5)
    r1 = p.add_run(bold_label + ' - ')
    set_run_font(r1, size=10.5, bold=True, color=BLACK)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=DARK_GREY)
    return p


# ---------------------------------------------------------------------------
# SUMMARY RENDERER
# ---------------------------------------------------------------------------

def _render_summary(doc, summary_text):
    """Split on double newline, single newline, or Seeking/Targeting keyword."""
    parts = [p.strip() for p in summary_text.split('\n\n') if p.strip()]

    if len(parts) == 1 and '\n' in summary_text:
        potential = [p.strip() for p in summary_text.split('\n') if p.strip()]
        if len(potential) > 1:
            parts = potential

    if len(parts) == 1:
        seeking_match = re.search(
            r'(?<=[.!?])\s+((?:Now\s+)?(?:[Ss]eeking|[Ll]ooking\s+for|[Tt]argeting\s+a|[Oo]pen\s+to)\b)',
            summary_text
        )
        if seeking_match:
            parts = [
                summary_text[:seeking_match.start()].strip(),
                summary_text[seeking_match.start():].strip()
            ]

    for i, part in enumerate(parts):
        if not part:
            continue
        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(part)
        set_run_font(r_sum, size=10.5)
        p_sum.paragraph_format.space_before = Pt(0)
        p_sum.paragraph_format.space_after  = Pt(6) if i == len(parts) - 1 else Pt(4)


# ---------------------------------------------------------------------------
# CV CONTENT BLOCK
# ---------------------------------------------------------------------------

def _add_cv_content(doc, cv_data):
    """Clean client-facing CV — name block through Technical Skills."""

    # Name — 26pt, bold, black, centered
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(cv_data['name'])
    set_run_font(r_name, size=26, bold=True, color=BLACK)
    p_name.paragraph_format.space_before = Pt(10)
    p_name.paragraph_format.space_after  = Pt(3)

    # Tagline — charcoal (not blue), centered
    p_titles = doc.add_paragraph()
    p_titles.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_titles = p_titles.add_run(cv_data['tagline'])
    set_run_font(r_titles, size=11, color=(51, 51, 51))
    p_titles.paragraph_format.space_after = Pt(2)

    # Contact — mid-grey, centered
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_contact = p_contact.add_run(cv_data['contact'])
    set_run_font(r_contact, size=10, color=(100, 100, 100))
    p_contact.paragraph_format.space_after = Pt(10)

    # Professional Summary
    add_cv_section_heading(doc, 'PROFESSIONAL SUMMARY')
    _render_summary(doc, cv_data.get('summary', ''))

    is_tech_role = cv_data.get('tech_role', False)

    # Technical Skills — top for tech roles
    if is_tech_role and cv_data.get('technical_skills'):
        add_cv_section_heading(doc, 'TECHNICAL SKILLS')
        p_tech = doc.add_paragraph()
        r_tech = p_tech.add_run(cv_data['technical_skills'])
        set_run_font(r_tech, size=10.5)

    # Skills / Competencies
    add_cv_section_heading(doc, 'SKILLS')
    p_comp = doc.add_paragraph()
    p_comp.paragraph_format.space_before = Pt(4)
    p_comp.paragraph_format.space_after  = Pt(6)
    r_comp = p_comp.add_run(cv_data['competencies'])
    set_run_font(r_comp, size=10.5, color=CORAL)

    # Work Experience
    add_cv_section_heading(doc, 'WORK EXPERIENCE')
    for role in cv_data['employment']:
        p_role = doc.add_paragraph()
        p_role.paragraph_format.space_before = Pt(8)
        p_role.paragraph_format.space_after  = Pt(1)
        r_role = p_role.add_run(role['header'])
        set_run_font(r_role, size=10.5, bold=True)

        if role.get('context'):
            p_ctx = doc.add_paragraph()
            p_ctx.paragraph_format.space_before = Pt(1)
            p_ctx.paragraph_format.space_after  = Pt(3)
            r_ctx = p_ctx.add_run(role['context'])
            set_run_font(r_ctx, size=10, italic=True, color=MID_GREY)

        for bullet in (role.get('bullets') or []):
            add_bullet(doc, bullet)

    # Additional Experience
    if cv_data.get('earlier_career'):
        add_cv_section_heading(doc, 'ADDITIONAL EXPERIENCE')
        for item in cv_data['earlier_career']:
            if isinstance(item, dict):
                title = item.get('title', '') or ''
                text  = item.get('text', '')  or ''
                line  = title + (' - ' + text if text else '')
            else:
                line = str(item) if item is not None else ''
            if line:
                add_bullet(doc, line)

    # Education
    if cv_data.get('education'):
        add_cv_section_heading(doc, 'EDUCATION')
        for item in cv_data['education']:
            add_bullet(doc, item)

    # Technical Skills — bottom for non-tech roles
    if not is_tech_role and cv_data.get('technical_skills'):
        add_cv_section_heading(doc, 'TECHNICAL SKILLS')
        p_tech = doc.add_paragraph()
        r_tech = p_tech.add_run(cv_data['technical_skills'])
        set_run_font(r_tech, size=10.5)


# ---------------------------------------------------------------------------
# REPORT CONTENT BLOCK
# ---------------------------------------------------------------------------

def _add_report_content(doc, cv_data):
    """Full DHM-branded report — foreword through sign-off."""

    first = _first_name(cv_data['name'])
    client_name = cv_data['name'].title()

    # --- Logo ---
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo.paragraph_format.space_before = Pt(0)
    p_logo.paragraph_format.space_after  = Pt(0)
    p_logo.add_run().add_picture(LOGO_PATH, width=Inches(1.8))

    p_logo_rule = doc.add_paragraph()
    p_logo_rule.paragraph_format.space_before = Pt(0)
    p_logo_rule.paragraph_format.space_after  = Pt(12)
    add_horizontal_rule(p_logo_rule, color='DC6A63')

    # --- Title ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after  = Pt(12)
    r_name = p_title.add_run(client_name)
    set_run_font(r_name, name='Inter', size=14, bold=True, color=CORAL)
    r_sub = p_title.add_run(' - CV Optimisation Report')
    set_run_font(r_sub, name='Inter', size=14, bold=True, color=DARK_GREY)

    p_title_rule = doc.add_paragraph()
    p_title_rule.paragraph_format.space_before = Pt(0)
    p_title_rule.paragraph_format.space_after  = Pt(18)
    add_horizontal_rule(p_title_rule, color='DC6A63')

    # --- Foreword ---
    p_lead = doc.add_paragraph()
    r_lead = p_lead.add_run('Your optimised CV is ready.')
    set_run_font(r_lead, size=10.5, color=CORAL)
    p_lead.paragraph_format.space_after = Pt(8)

    add_body(doc,
        'It has been completely rebuilt - the structure, the language, the framing of your '
        'achievements, and the keywords that will carry you through automated screening before '
        'a recruiter ever sees your name.'
    )
    add_body(doc,
        'You have done the hard work to get here. This document finally reflects it properly.'
    )

    # --- How We Approach CV Optimisation ---
    add_report_section_heading(doc, 'How We Approach CV Optimisation')
    add_subheading(doc, 'What we consider, and why it matters.')

    add_body(doc,
        'Every CV we optimise is evaluated against two distinct stages of the hiring process. '
        'Both matter. Failing either one ends your application before it has begun.'
    )

    add_stage_heading(doc, 'Stage One: The Technology')
    add_body(doc,
        'Before a recruiter reads a single word, your CV is processed by an Applicant Tracking '
        'System. ATS software parses your document and scores it against the job specification. '
        'If the right signals are not present, the CV is rejected automatically. No human sees it.'
    )
    add_body(doc,
        'Those signals include formatting, keywords, job title alignment, section headings, '
        'achievement framing, and industry-specific terminology. Each of these has been addressed '
        'in your optimised CV, because getting any one of them wrong is enough to trigger an '
        'automated rejection.'
    )
    add_body(doc,
        'Your CV has been saved as a Word document (.docx), the standard format for ATS '
        'submissions. PDF and PowerPoint files are not reliably parsed by all systems and can '
        'result in an automatic rejection. We recommend keeping it in this format when applying.'
    )

    add_stage_heading(doc, 'Stage Two: The Human')
    add_body(doc,
        'If your CV clears the ATS stage, it reaches a recruiter or hiring manager. At this '
        'point, the challenge changes. Research consistently shows that initial CV decisions are '
        'made within the first few seconds of opening a document. In that window, the reader is '
        'scanning for one thing: evidence that this person can do what we need.'
    )
    add_body(doc,
        'To make sure your CV holds attention beyond that first scan, we focus on several things:',
        space_after=4
    )

    approach_bullets = [
        ('Executive Summary',
         'this is the first thing a hiring manager reads. Yours has been written to lead with '
         'impact, drawing out your strongest achievements and most relevant experience so the '
         'reader knows immediately what you bring to the table.'),
        ('Achievement framing',
         'every role focuses on results and the impact you had on the business, not just '
         'responsibilities. Hiring managers are looking for evidence of contribution, not a '
         'list of duties.'),
        ('Career progression clarity',
         'your career story is presented in a way that is easy to follow, demonstrating growth '
         'and increasing responsibility where it exists.'),
        ('Relevance to the target role',
         'the language, experience, and skills most relevant to the roles you are targeting have '
         'been brought forward. What is less relevant has been deprioritised.'),
        ('Consistency of tone and language',
         'the document reads as a coherent whole, with a professional and consistent voice '
         'throughout.'),
        ('Removal of anything working against you',
         'anything that could create an unconscious bias or distract from your strengths has '
         'been reviewed and addressed.'),
    ]
    for label, text in approach_bullets:
        add_approach_bullet(doc, label, text)

    add_stage_heading(doc, 'Why both stages matter')
    add_body(doc,
        'A CV that passes ATS but reads poorly to a human does not get an interview. A CV that '
        'reads well but fails ATS never reaches a human at all. Every decision made in your '
        'optimised CV - formatting, structure, language, and framing - has been made with both '
        'stages in mind.'
    )
    add_body(doc,
        'The goal is straightforward: to get you in front of the right person, face to face, '
        'where you can do what no CV can do. Bring your story to life.',
        bold=True,
        space_after=10
    )

    # --- What I Changed ---
    add_report_section_heading(doc, 'What I Changed - and Why')
    add_subheading(doc, 'Every change was deliberate. Here is the thinking behind each one.')

    for i, item in enumerate(cv_data.get('changelog', []), 1):
        add_numbered_item(doc, i, item['title'], item['text'])

    # --- My Honest Assessment ---
    doc.add_paragraph()
    add_report_section_heading(doc, 'My Honest Assessment')
    add_subheading(doc,
        f'The CV is stronger, {first}. These are the things worth knowing as you head into your search.'
    )

    for i, item in enumerate(cv_data.get('gap_report', []), 1):
        add_numbered_item(doc, i, item['title'], item['text'])

    p_gap_note = doc.add_paragraph()
    r_gap_note = p_gap_note.add_run(
        'Addressing even one or two of these points before you start applying will put you '
        'meaningfully ahead. The changes are straightforward. The difference they make to your '
        'results will not be.'
    )
    set_run_font(r_gap_note, size=10.5, bold=True, color=BLACK)
    p_gap_note.paragraph_format.space_before = Pt(6)
    p_gap_note.paragraph_format.space_after  = Pt(24)

    # --- Thank You ---
    add_report_section_heading(doc, 'Thank You')
    add_subheading(doc, 'It is a privilege to work on something that matters this much.')

    add_body(doc,
        'Thank you for choosing Dear Hiring Manager. Helping people navigate a system that '
        'was not built in their favour is why we do this. We hope your optimised CV opens '
        'the doors it deserves to.'
    )
    add_body(doc,
        'When you are ready to go further, we can help with three things:',
        space_after=4
    )

    next_steps = [
        ('LinkedIn Profile Optimisation',
         'aligning your LinkedIn presence with your optimised CV so every touchpoint tells '
         'the same story. Recruiters check both. They need to match.'),
        ('Job Search Strategy',
         'a structured approach to targeting the right roles, in the right organisations, '
         'in the right way. More signal. Less noise.'),
        ('Interview Preparation and Mock Interviews',
         'when the CV gets you through the door, you need to be ready for what comes next. '
         'We work with you to prepare for interviews, sharpen your answers, and walk in '
         'with confidence.'),
    ]
    for label, text in next_steps:
        add_next_step_bullet(doc, label, text)

    p_cta = doc.add_paragraph()
    p_cta.paragraph_format.space_before = Pt(6)
    p_cta.paragraph_format.space_after  = Pt(6)
    r_cta_pre = p_cta.add_run('All three are available at ')
    set_run_font(r_cta_pre, size=10.5, bold=True, color=CORAL)
    add_hyperlink(p_cta, 'https://www.dearhiringmanager.careers',
                  'www.DearHiringManager.careers', CORAL, size=10.5, bold=True)
    r_cta_dot = p_cta.add_run('.')
    set_run_font(r_cta_dot, size=10.5, bold=True, color=CORAL)

    # --- Sign-off ---
    doc.add_paragraph().paragraph_format.space_after = Pt(8)
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_sig.paragraph_format.space_before = Pt(0)
    p_sig.paragraph_format.space_after  = Pt(0)
    p_sig.add_run().add_picture(SIG_PATH, width=Inches(4.25))


# ---------------------------------------------------------------------------
# PUBLIC BUILD FUNCTIONS
# ---------------------------------------------------------------------------

def build_cv_only(cv_data, output_path):
    """Client-facing CV — clean, no DHM branding."""
    doc = Document()
    _setup_page(doc)
    _add_cv_content(doc, cv_data)
    doc.save(output_path)
    print(f'Saved CV: {output_path}')


def build_report_only(cv_data, output_path):
    """DHM-branded strategic report — foreword through sign-off."""
    doc = Document()
    _setup_page(doc)
    _add_report_content(doc, cv_data)
    doc.save(output_path)
    print(f'Saved Report: {output_path}')


def build_cv_doc(cv_data, output_path):
    """Combined document — CV + report (legacy / testing use)."""
    doc = Document()
    _setup_page(doc)

    p_s1 = doc.add_paragraph()
    r_s1 = p_s1.add_run('SECTION 1 - REWRITTEN CV DRAFT')
    set_run_font(r_s1, name='Inter', size=11, bold=True, color=CORAL)
    r_s1.font.underline = True
    p_s1.paragraph_format.space_after = Pt(10)

    _add_cv_content(doc, cv_data)

    p_break = doc.add_paragraph()
    run_break = p_break.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run_break._r.append(br)

    _add_report_content(doc, cv_data)
    doc.save(output_path)
    print(f'Saved combined: {output_path}')


# ---------------------------------------------------------------------------
# TEST DATA
# ---------------------------------------------------------------------------

if __name__ == '__main__':
    cv_data = {
        'tech_role': False,
        'name': 'JAMES CARTER',
        'tagline': 'Senior Marketing Manager | Head of Marketing | Demand Generation Lead',
        'contact': '07XXX XXXXXX  •  james.carter@email.com  •  London, UK',
        'summary': (
            'A demand generation leader who built TechFlow\'s entire marketing function from '
            'scratch, generating £4.2M in attributed pipeline revenue - 67% year-on-year growth - '
            'and cutting customer acquisition cost by 31% in the process. Known for building '
            'data-driven marketing engines that align tightly with sales, with a consistent '
            'focus on pipeline impact over brand activity.\n\n'
            'Now seeking a Head of Marketing or Senior Marketing Manager role with full '
            'ownership of pipeline strategy, team, and budget.'
        ),
        'competencies': (
            'Demand Generation Strategy  •  B2B SaaS Marketing  •  Account-Based Marketing (ABM)  '
            '•  Paid Search and Paid Media  •  SEO and Content Strategy  •  Budget Ownership and '
            'ROI Reporting  •  CRM and Marketing Automation (HubSpot, Salesforce)  •  Marketing '
            'Attribution  •  Team Leadership  •  Stakeholder Management'
        ),
        'employment': [
            {
                'header': 'MARKETING MANAGER  |  TECHFLOW LTD  |  Feb 2021 - Present',
                'context': 'Series A B2B SaaS platform. First marketing hire; built the function from scratch.',
                'bullets': [
                    'Generated £4.2M in attributed pipeline revenue in FY2024 - 67% YoY - recognised by the CEO as TechFlow\'s single biggest commercial growth driver.',
                    'Cut wasted paid search spend from £12K to £4K per month and reduced customer acquisition cost by 31% while growing pipeline volume by 67% YoY.',
                    'Built a 120-article SEO library, scaling organic traffic from 8,000 to 47,000 monthly sessions in 18 months.',
                ]
            }
        ],
        'earlier_career': [
            'Marketing Assistant  |  Digital Spark Agency  |  Sep 2016 - Feb 2018  -  '
            'Grew combined client social following by 22,000 in 12 months; built email nurture '
            'sequences achieving 38% open rate against a 21% industry benchmark.'
        ],
        'education': [
            'BA (Hons) Marketing, 2:1  -  University of Leeds (2013-2016)',
            'HubSpot Marketing Software Certification  -  2023',
        ],
        'technical_skills': (
            'HubSpot  •  Salesforce  •  Google Analytics 4  •  SEMrush  •  Google Ads  '
            '•  Meta Ads  •  LinkedIn Campaign Manager  •  Looker Studio'
        ),
        'changelog': [
            {'title': 'ATS title injection',
             'text': 'Your target job titles now appear directly under your name. ATS systems match on exact title strings before a human reads a word. Candidates without those strings are filtered out regardless of their experience.'},
            {'title': 'Summary rebuilt around your strongest result',
             'text': 'The original CV opened with context. Your optimised version opens with the number that matters most - £4.2M attributed pipeline, 67% year-on-year. That is what earns the next six seconds of attention.'},
            {'title': 'Skills section built from scratch',
             'text': 'This section did not exist in your original CV. A keyword-dense skills block gives the ATS a clean match signal and gives every human reader an immediate picture of what you bring.'},
        ],
        'gap_report': [
            {'title': 'Team leadership evidence needs strengthening [HIGH PRIORITY]',
             'text': 'You led a team of two, and the CV carries almost no evidence of how. At Head of Marketing level, people leadership is a primary screening criterion. One specific achievement bullet per role on coaching, hiring, or development outcomes will close this gap significantly.'},
            {'title': 'Marketing budget ownership is not stated [MEDIUM PRIORITY]',
             'text': 'The annual budget you owned at TechFlow does not appear anywhere. Budget ownership is a standard filter at Senior Manager level and above. Adding the figure changes the weight of the role immediately.'},
        ],
    }

    build_cv_only(cv_data, '/sessions/kind-epic-bardeen/mnt/outputs/DHM_CV_James_Carter.docx')
    build_report_only(cv_data, '/sessions/kind-epic-bardeen/mnt/outputs/DHM_Report_James_Carter.docx')
